import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def summary_Freq(data, activity_codes, save_location):
    """Creates bar chart of activity frequencies."""
    Activity = []
    Counts = []
    for key in activity_codes.keys():
        Activity.append(key)
        Counts.append(data.pathways.str.contains(key).sum())

    ActivityDF = pd.DataFrame({
        'Activity': Activity,
        'Count': Counts
    })

    figure  = plt.Figure(figsize=(5, 5))
    ax =  figure.add_subplot(111)
    ax.barh(range(len(activity_codes)), ActivityDF['Count'],color='#A60628')   
    
    ax.set_xlabel('Frequency')
    ax.set_title('Frequency Occurance Per Activity')
    ax.set_yticks(range(len(activity_codes)))
    ax.set_yticklabels(ActivityDF['Activity'], wrap=True, va='center')
    for patch in ax.patches:
        bl = patch.get_xy()
        x = 0.5 * patch.get_width() + bl[0]
        y = 0.3 * patch.get_height() + bl[1] 
        ax.text(x,y,"%d" %(patch.get_width()),
                ha='center', color ='white', weight='bold', size='12')

    file_name = save_location + 'Plots/Summary/Activity_Frequency.png'
    figure.savefig(file_name, bbox_inches='tight', facecolor="None")    
    plt.close()
                

def histAll(data, save_location):
    """Creates histogram of all total times in system."""
    figure  = plt.Figure(figsize=(5, 5))
    ax =  figure.add_subplot(111)
    data.hist(column='totaltime', bins=30, ax=ax)
    ax.set_xlabel('Number of days')
    ax.set_ylabel('Frequency')
    ax.set_title('Total Time of Pathway')

    file_name = save_location + 'Plots/Summary/Histogram_Total_Time_in_System.png'
    figure.savefig(file_name,bbox_inches='tight', facecolor="None")      
    plt.close()


def boxplotAll(data, activity_codes, save_location, original_name):
    """Creates boxplot of waiting time for all activities."""
    figure, ax  = plt.subplots(1, 1, figsize=(5, 5))
    
    plot_dict = {}
    for letter in activity_codes.keys():
        waits = [data[letter][i] for i in range(len(data))]
        if original_name == 'original':
            waits = [x for x in waits if str(x) != 'nan']
        elif original_name == 'original_formatted':
            waits = [y for x in waits for y in x if str(y) != 'nan']
        plot_dict[letter] = waits
        
    ax.boxplot(plot_dict.values())
    ax.set_xticklabels(plot_dict.keys())
    ax.set_xlabel('Activity')
    ax.set_ylabel('Time (days)')
    ax.set_title('Boxplot of Wait Time to Activity')

    file_name = save_location + 'Plots/Summary/Boxplot_Activity_Wait_Time.png'
    figure.savefig(file_name,bbox_inches='tight', facecolor="None")     
    plt.close()


def heatAll(data, activity_codes, save_location):
    """Creates a heatplot of all unique pathways."""
    max_length = len(max(data['pathways'], key=len))
    all_pathways =[]

    for row in data.pathways:
        all_pathways.append(list(row.ljust(max_length,' ')))

    letters = [code for code in activity_codes.keys()]
    letters.insert(0,' ')

    converted_pathways = []
    for row in all_pathways:
        convert_letters = [letters.index(r) - 1 for r in row]
        converted_pathways.append(convert_letters)

    fig, ax = plt.subplots(1, figsize=(5, 6))

    hm = ax.pcolor(converted_pathways, cmap ="cubehelix_r", vmin=0, vmax=len(letters)-1)
    hm.cmap.set_under('white')
    ax.set_xticks([i+0.5 for i in range(int(max_length))])
    ax.set_xticklabels([i+1 for i in range(int(max_length))],ha='center')
    ax.set_title('Heatmap of All Data', size=12)

    cbar = plt.colorbar(hm)
    cbar.set_ticks([i + 0.5 for i in range(len(activity_codes))])
    cbar.set_ticklabels(letters[1:])

    plt.ylabel('Patient Pathway Number', size=12)
    plt.xlabel('Activity Position', size=12)

    file_name = save_location + 'Plots/Summary/Heatmap_All_Pathways.png'
    fig.savefig(file_name,bbox_inches='tight', facecolor="None")  
    
    
def SummarySheet(data, df, activity_codes, save_location, original_name):
    """Creates a word document with summary infomration about the data."""
    num_patients = str(data['pathways'].count())
    num_activitites = str(len(activity_codes))
    num_pathways = str(len(data.pathways.value_counts()))
    num_pathways_ex1 = str(len(data['pathways'].value_counts()[data['pathways'].value_counts()<2]))
    summary_Freq(data, activity_codes, save_location)
    index_occur = [pathway for i, pathway in enumerate(df.pathway) if i < 10]
    num_occur = [count for i, count in enumerate(df.counts) if i < 10]
    heatAll(data, activity_codes, save_location)
    mean_total = str(round(data.totaltime.mean(),2))
    median_total = str(round(data.totaltime.median(),2))
    per25_total = str(round(data.totaltime.quantile(0.25),2))
    per75_total = str(round(data.totaltime.quantile(0.75),2))
    histAll(data, save_location)
    boxplotAll(data, activity_codes, save_location, original_name)

    # Create Document
    document = Document()
    document.add_heading('Summary Sheet', 0)
    
    document.add_heading('Data Summary', 1)
    paragraph = document.add_paragraph('The data includes ' + num_patients + ' patiens, ' + num_activitites + ' activities and ' + num_pathways + ' different pathways, where ' + num_pathways_ex1 + ' pathways only occured once.'
                                   ' For reference, Table 1 shows a key of activities and their codes.')
    table_letters = document.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells = table_letters.rows[0].cells
    hdr_cells[0].text = 'Code'
    hdr_cells[1].text = 'Activity'
    for key, value in activity_codes.items():
        row_cells = table_letters.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value

    document.add_heading('Activity Summary', 1)
    document.add_paragraph('The frequency of occurrence for each activity can be seen in Figure 1.')
    document.add_picture(save_location + 'Plots/Summary/Activity_Frequency.png', width=Inches(4))
    
    document.add_heading('Pathway Summary', 1)
    table_occur = document.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells = table_occur.rows[0].cells
    hdr_cells[0].text = 'Pathway'
    hdr_cells[1].text = 'Frequency'
    maximum = 10
    if int(num_pathways) < 10:
        maximum = int(num_pathways)
    document.add_paragraph('The ' + str(maximum) + ' most popular pathways are:')
    for i in range(maximum):
        row_cells = table_occur.add_row().cells
        row_cells[0].text = index_occur[i]
        row_cells[1].text = str(num_occur[i])
    document.add_paragraph('\n The heatmap in Figure 2 is a visual representation of the all the pathways included in the data.')
    document.add_picture(save_location + 'Plots/Summary/Heatmap_All_Pathways.png', width=Inches(6))

    document.add_heading('Time Summary', 1)
    document.add_paragraph('The mean, median, 25 percentile and 75 percentile total time was ' + mean_total + ', ' + median_total + ', ' + per25_total + ' and ' + per75_total + ' days respectively. The histogram in Figure 3 displays the overall total times.')
    document.add_picture(save_location + 'Plots/Summary/Histogram_Total_Time_in_System.png', width=Inches(4))
    document.add_paragraph('The time to each activity from the one that preceeded it is displayed in the boxplot in Figure 4')
    document.add_picture(save_location + 'Plots/Summary/Boxplot_Activity_Wait_Time.png', width=Inches(4))
    file_name = save_location + 'Summary_Sheet.docx'
    document.save(file_name)